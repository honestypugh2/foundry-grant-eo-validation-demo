"""
Document Ingestion Agent
Handles file upload, OCR, and metadata extraction from grant proposals.
"""

import os
import logging
from typing import Dict, Any, Optional
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentIngestionAgent:
    """
    Agent responsible for ingesting documents, extracting text and metadata.
    Supports PDF, DOCX, and TXT files with OCR capabilities via Azure Document Intelligence.
    """
    
    def __init__(self, use_azure: bool = False):
        """
        Initialize the Document Ingestion Agent.
        
        Args:
            use_azure: If True, use Azure Document Intelligence for OCR. Otherwise, use local processing.
        """
        self.use_azure = use_azure
        self.azure_endpoint = os.getenv('AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT')
        self.azure_key = os.getenv('AZURE_DOCUMENT_INTELLIGENCE_API_KEY')
        
        if self.use_azure and not self.azure_endpoint:
            logger.warning("Azure Document Intelligence not configured. Falling back to local processing.")
            self.use_azure = False
    
    def process_document(self, file_path: str, file_name: Optional[str] = None) -> Dict[str, Any]:
        """
        Process a document and extract text and metadata.
        
        Args:
            file_path: Path to the document file
            file_name: Optional custom filename
            
        Returns:
            Dictionary containing extracted text, metadata, and processing info
        """
        logger.info(f"Processing document: {file_path}")
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = Path(file_path).suffix.lower()
        actual_file_name = file_name or os.path.basename(file_path)
        
        try:
            if self.use_azure:
                result = self._process_with_azure(file_path)
            else:
                result = self._process_locally(file_path, file_extension)
            
            # Add metadata
            result['metadata'] = {
                'filename': actual_file_name,
                'file_path': file_path,
                'file_size': os.path.getsize(file_path),
                'file_type': file_extension,
                'processed_at': datetime.now().isoformat(),
                'processing_method': 'azure' if self.use_azure else 'local'
            }
            
            logger.info(f"Successfully processed document: {actual_file_name}")
            return result
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise
    
    def _process_locally(self, file_path: str, file_extension: str) -> Dict[str, Any]:
        """Process document using local methods (no OCR)."""
        text = ""
        
        if file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        
        elif file_extension == '.pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            except ImportError:
                logger.warning("PyPDF2 not installed. PDF text extraction limited.")
                text = f"[PDF file: {os.path.basename(file_path)}. Install PyPDF2 for text extraction.]"
        
        elif file_extension == '.docx':
            try:
                import docx
                doc = docx.Document(file_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            except ImportError:
                logger.warning("python-docx not installed. DOCX text extraction limited.")
                text = f"[DOCX file: {os.path.basename(file_path)}. Install python-docx for text extraction.]"
        
        else:
            text = f"[Unsupported file type: {file_extension}]"
        
        return {
            'text': text,
            'page_count': text.count('\n') // 50 + 1,  # Rough estimate
            'word_count': len(text.split()),
            'char_count': len(text)
        }
    
    def _process_with_azure(self, file_path: str) -> Dict[str, Any]:
        """Process document using Azure Document Intelligence (OCR-enabled)."""
        try:
            from azure.ai.formrecognizer import DocumentAnalysisClient
            from azure.core.credentials import AzureKeyCredential
            
            if not self.azure_endpoint or not self.azure_key:
                raise ValueError("Azure Document Intelligence endpoint and key are required")
            
            client = DocumentAnalysisClient(
                endpoint=self.azure_endpoint,
                credential=AzureKeyCredential(self.azure_key)
            )
            
            with open(file_path, 'rb') as f:
                poller = client.begin_analyze_document("prebuilt-document", document=f)
                result = poller.result()
            
            # Extract text
            text = result.content
            
            # Extract additional metadata
            tables = []
            if result.tables:
                for table in result.tables:
                    table_data = {
                        'row_count': table.row_count,
                        'column_count': table.column_count,
                        'cells': [{'content': cell.content, 'row': cell.row_index, 'col': cell.column_index} 
                                 for cell in table.cells]
                    }
                    tables.append(table_data)
            
            key_value_pairs = []
            if result.key_value_pairs:
                for kv_pair in result.key_value_pairs:
                    if kv_pair.key and kv_pair.value:
                        key_value_pairs.append({
                            'key': kv_pair.key.content,
                            'value': kv_pair.value.content
                        })
            
            return {
                'text': text,
                'page_count': len(result.pages),
                'word_count': len(text.split()),
                'char_count': len(text),
                'tables': tables,
                'key_value_pairs': key_value_pairs
            }
            
        except Exception as e:
            logger.warning(f"Azure processing failed: {str(e)}. Falling back to local processing.")
            return self._process_locally(file_path, Path(file_path).suffix.lower())
    
    def extract_metadata(self, document_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract key metadata fields from processed document.
        
        Args:
            document_data: Processed document data
            
        Returns:
            Dictionary of metadata fields
        """
        metadata = document_data.get('metadata', {})
        
        # Extract common fields from grant proposals
        extracted = {
            'document_type': 'grant_proposal',
            'word_count': document_data.get('word_count', 0),
            'page_count': document_data.get('page_count', 0),
            'processing_timestamp': metadata.get('processed_at'),
            'file_name': metadata.get('filename'),
            'file_size': metadata.get('file_size')
        }
        
        # Try to extract key-value pairs if available (from Azure processing)
        if 'key_value_pairs' in document_data:
            for kv in document_data['key_value_pairs']:
                key = kv['key'].lower()
                if 'date' in key or 'deadline' in key:
                    extracted['deadline'] = kv['value']
                elif 'amount' in key or 'budget' in key:
                    extracted['budget_amount'] = kv['value']
                elif 'organization' in key or 'applicant' in key:
                    extracted['applicant'] = kv['value']
        
        return extracted
